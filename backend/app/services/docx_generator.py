"""
DOCX Resume Generator
======================
Takes an optimized resume JSON and generates a professional,
ATS-friendly DOCX document using python-docx.

Design principles:
- NO tables, text boxes, or graphics (ATS-safe)
- Standard section headers (EXPERIENCE, EDUCATION, SKILLS, …)
- Clean typography: Calibri, consistent spacing
- Machine-readable date formats
- Single-column layout
"""

from typing import Dict, Any, List, Optional
from io import BytesIO
import logging

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

# ── Style constants ───────────────────────────────────────────────────
FONT_NAME = "Calibri"
FONT_SIZE_NAME = Pt(20)
FONT_SIZE_CONTACT = Pt(10)
FONT_SIZE_SECTION_HEADER = Pt(12)
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_SMALL = Pt(9.5)
COLOR_PRIMARY = RGBColor(0x1A, 0x1A, 0x1A)      # near-black
COLOR_SECONDARY = RGBColor(0x55, 0x55, 0x55)     # dark gray
COLOR_ACCENT = RGBColor(0x1A, 0x56, 0xA8)        # blue for links/name
SECTION_SPACING_BEFORE = Pt(10)
SECTION_SPACING_AFTER = Pt(4)
BULLET_CHAR = "\u2022"  # •


class DocxGenerator:
    """Generate an ATS-friendly DOCX from optimized resume JSON."""

    def generate(self, resume_data: Dict[str, Any]) -> bytes:
        """
        Generate a DOCX file from resume JSON.

        IMPORTANT: This intentionally only reads the candidate-facing
        resume fields (name, contact, summary, skills, experience,
        projects, education, certifications). Internal metadata such
        as `optimization_metadata` (which contains the gap_analysis
        and ATS deltas) MUST NEVER appear in the printed resume —
        that information is for the user's UI, not the employer.

        Args:
            resume_data: Optimized resume JSON (flat schema)

        Returns:
            DOCX file content as bytes
        """
        doc = Document()
        self._setup_styles(doc)
        self._set_margins(doc)

        # ── Header: Name + Contact ────────────────────────────────────
        self._add_name(doc, resume_data.get("name", ""))
        self._add_contact_line(doc, resume_data)

        # ── Summary ───────────────────────────────────────────────────
        summary = resume_data.get("summary")
        if summary:
            self._add_section_header(doc, "PROFESSIONAL SUMMARY")
            self._add_paragraph(doc, summary)

        # ── Skills ────────────────────────────────────────────────────
        skills = resume_data.get("skills", [])
        if skills:
            self._add_section_header(doc, "SKILLS")
            self._add_skills(doc, skills)

        # ── Experience ────────────────────────────────────────────────
        experience = resume_data.get("experience", [])
        if experience:
            self._add_section_header(doc, "EXPERIENCE")
            for exp in experience:
                self._add_experience_entry(doc, exp)

        # ── Projects ──────────────────────────────────────────────────
        projects = resume_data.get("projects", [])
        if projects:
            self._add_section_header(doc, "PROJECTS")
            for proj in projects:
                self._add_project_entry(doc, proj)

        # ── Education ─────────────────────────────────────────────────
        education = resume_data.get("education", [])
        if education:
            self._add_section_header(doc, "EDUCATION")
            for edu in education:
                self._add_education_entry(doc, edu)

        # ── Certifications ────────────────────────────────────────────
        certs = resume_data.get("certifications", [])
        if certs:
            self._add_section_header(doc, "CERTIFICATIONS")
            for cert in certs:
                self._add_bullet(doc, cert)

        # ── Serialize to bytes ────────────────────────────────────────
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()
        logger.info(f"Generated DOCX: {len(docx_bytes)} bytes")
        return docx_bytes

    # ── Style setup ───────────────────────────────────────────────────

    @staticmethod
    def _setup_styles(doc: Document):
        """Configure document-level default styles."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = FONT_NAME
        font.size = FONT_SIZE_BODY
        font.color.rgb = COLOR_PRIMARY
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = 1.15

    @staticmethod
    def _set_margins(doc: Document):
        """Set narrow margins for more content space (ATS-safe)."""
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

    # ── Header components ─────────────────────────────────────────────

    def _add_name(self, doc: Document, name: str):
        """Add candidate name as document title."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name or "")
        run.bold = True
        run.font.size = FONT_SIZE_NAME
        run.font.color.rgb = COLOR_ACCENT
        run.font.name = FONT_NAME
        p.paragraph_format.space_after = Pt(2)

    def _add_contact_line(self, doc: Document, data: Dict[str, Any]):
        """Add a single centered contact line: email | phone | linkedin | github | location."""
        parts = []
        for field in ("email", "phone", "linkedin", "github", "portfolio", "location"):
            val = data.get(field)
            if val:
                # Defensive: LLM sometimes returns lists for contact fields
                if isinstance(val, list):
                    val = val[0] if val else None
                if val:
                    parts.append(str(val))

        if not parts:
            return

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("  |  ".join(parts))
        run.font.size = FONT_SIZE_CONTACT
        run.font.color.rgb = COLOR_SECONDARY
        run.font.name = FONT_NAME
        p.paragraph_format.space_after = Pt(6)

    # ── Section header ────────────────────────────────────────────────

    def _add_section_header(self, doc: Document, title: str):
        """Add a bold section header with bottom border."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = SECTION_SPACING_BEFORE
        p.paragraph_format.space_after = SECTION_SPACING_AFTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = FONT_SIZE_SECTION_HEADER
        run.font.color.rgb = COLOR_PRIMARY
        run.font.name = FONT_NAME

        # Bottom border via XML (thin line under section header)
        from docx.oxml.ns import qn
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Content helpers ───────────────────────────────────────────────

    def _add_paragraph(self, doc: Document, text: str, bold: bool = False):
        """Add a plain paragraph."""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = FONT_SIZE_BODY
        run.font.name = FONT_NAME
        run.bold = bold

    def _add_skills(self, doc: Document, skills: List[str]):
        """Add skills as a comma-separated paragraph (ATS-readable)."""
        # Defensive: flatten any nested lists and coerce to strings
        flat = []
        for s in skills:
            if isinstance(s, list):
                flat.extend(str(item) for item in s)
            else:
                flat.append(str(s))
        p = doc.add_paragraph()
        run = p.add_run(", ".join(flat))
        run.font.size = FONT_SIZE_BODY
        run.font.name = FONT_NAME

    def _add_experience_entry(self, doc: Document, exp: Dict[str, Any]):
        """Add a single experience block: title | company | dates + bullets."""
        # Title line: "Software Engineer | Company Inc."
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

        title = exp.get("title", "")
        company = exp.get("company", "")

        run_title = p.add_run(title)
        run_title.bold = True
        run_title.font.size = FONT_SIZE_BODY
        run_title.font.name = FONT_NAME

        if company:
            run_sep = p.add_run("  |  ")
            run_sep.font.size = FONT_SIZE_BODY
            run_sep.font.color.rgb = COLOR_SECONDARY
            run_sep.font.name = FONT_NAME

            run_company = p.add_run(company)
            run_company.font.size = FONT_SIZE_BODY
            run_company.font.color.rgb = COLOR_SECONDARY
            run_company.font.name = FONT_NAME

        # Dates line
        dates = exp.get("dates", "")
        if dates:
            p_dates = doc.add_paragraph()
            p_dates.paragraph_format.space_before = Pt(0)
            p_dates.paragraph_format.space_after = Pt(2)
            run_dates = p_dates.add_run(dates)
            run_dates.italic = True
            run_dates.font.size = FONT_SIZE_SMALL
            run_dates.font.color.rgb = COLOR_SECONDARY
            run_dates.font.name = FONT_NAME

        # Bullets
        for bullet in (exp.get("description") or []):
            self._add_bullet(doc, bullet)

    def _add_project_entry(self, doc: Document, proj: Dict[str, Any]):
        """Add a single project block."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(1)

        name = proj.get("name", "")
        run_name = p.add_run(name)
        run_name.bold = True
        run_name.font.size = FONT_SIZE_BODY
        run_name.font.name = FONT_NAME

        # Technologies
        techs = proj.get("technologies", [])
        if techs:
            run_tech = p.add_run(f"  —  {', '.join(techs)}")
            run_tech.font.size = FONT_SIZE_SMALL
            run_tech.font.color.rgb = COLOR_SECONDARY
            run_tech.font.name = FONT_NAME

        # Description
        desc = proj.get("description")
        if desc:
            p_desc = doc.add_paragraph()
            p_desc.paragraph_format.space_before = Pt(0)
            run_desc = p_desc.add_run(desc)
            run_desc.font.size = FONT_SIZE_BODY
            run_desc.font.name = FONT_NAME

        # Bullets
        for bullet in (proj.get("bullets") or []):
            self._add_bullet(doc, bullet)

        # Link
        link = proj.get("link")
        if link:
            p_link = doc.add_paragraph()
            p_link.paragraph_format.space_before = Pt(0)
            run_link = p_link.add_run(link)
            run_link.font.size = FONT_SIZE_SMALL
            run_link.font.color.rgb = COLOR_ACCENT
            run_link.font.name = FONT_NAME

    def _add_education_entry(self, doc: Document, edu: Dict[str, Any]):
        """Add a single education block."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)

        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        year = edu.get("year", "")

        run_degree = p.add_run(degree)
        run_degree.bold = True
        run_degree.font.size = FONT_SIZE_BODY
        run_degree.font.name = FONT_NAME

        if institution:
            run_sep = p.add_run("  |  ")
            run_sep.font.size = FONT_SIZE_BODY
            run_sep.font.color.rgb = COLOR_SECONDARY
            run_sep.font.name = FONT_NAME

            run_inst = p.add_run(institution)
            run_inst.font.size = FONT_SIZE_BODY
            run_inst.font.name = FONT_NAME

        if year:
            run_year = p.add_run(f"  ({year})")
            run_year.font.size = FONT_SIZE_SMALL
            run_year.font.color.rgb = COLOR_SECONDARY
            run_year.font.name = FONT_NAME

    def _add_bullet(self, doc: Document, text: str):
        """Add a bullet point line."""
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(f"{BULLET_CHAR}  {text}")
        run.font.size = FONT_SIZE_BODY
        run.font.name = FONT_NAME


# Singleton
docx_generator = DocxGenerator()
