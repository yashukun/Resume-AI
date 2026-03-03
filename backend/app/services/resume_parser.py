import fitz  # PyMuPDF
import pymupdf4llm
import mammoth
import io
import re
from typing import Dict, Any, List, Tuple
from bs4 import BeautifulSoup
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import logging

logger = logging.getLogger(__name__)

# Patterns for categorizing extracted URLs
LINK_PATTERNS = {
    "linkedin": re.compile(r"linkedin\.com", re.IGNORECASE),
    "github": re.compile(r"github\.com", re.IGNORECASE),
    "portfolio": re.compile(r"(portfolio|personal|\.dev|\.me|\.io)", re.IGNORECASE),
    "twitter": re.compile(r"(twitter\.com|x\.com)", re.IGNORECASE),
    "leetcode": re.compile(r"leetcode\.com", re.IGNORECASE),
    "medium": re.compile(r"medium\.com", re.IGNORECASE),
    "stackoverflow": re.compile(r"stackoverflow\.com", re.IGNORECASE),
    "kaggle": re.compile(r"kaggle\.com", re.IGNORECASE),
    "behance": re.compile(r"behance\.net", re.IGNORECASE),
    "dribbble": re.compile(r"dribbble\.com", re.IGNORECASE),
}

EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_PATTERN = re.compile(r"(\+?\d[\d\-\.\s\(\)]{7,}\d)")
URL_PATTERN = re.compile(r"https?://[^\s\)\]\},\"']+")


def _categorize_links(urls: List[str]) -> Dict[str, List[str]]:
    """
    Categorize a list of URLs into known platforms.

    Returns dict like {"linkedin": ["https://..."], "github": ["https://..."], "other": [...]}
    """
    categorized: Dict[str, List[str]] = {}
    for url in urls:
        url = url.strip().rstrip(".,;:)")
        if not url:
            continue
        matched = False
        for category, pattern in LINK_PATTERNS.items():
            if pattern.search(url):
                categorized.setdefault(category, []).append(url)
                matched = True
                break
        if not matched:
            categorized.setdefault("other", []).append(url)
    return categorized


def _extract_emails_from_text(text: str) -> List[str]:
    """Extract email addresses from text."""
    return list(set(EMAIL_PATTERN.findall(text)))


def _extract_phones_from_text(text: str) -> List[str]:
    """Extract phone numbers from text."""
    phones = []
    for match in PHONE_PATTERN.findall(text):
        cleaned = re.sub(r"[\s\-\.\(\)]", "", match)
        if len(cleaned) >= 8:  # Minimum viable phone number length
            phones.append(match.strip())
    return list(set(phones))


def _extract_urls_from_text(text: str) -> List[str]:
    """Extract URLs from plain text as fallback."""
    return list(set(URL_PATTERN.findall(text)))


class ResumeParserService:
    """
    Hybrid resume parser that extracts structure, text, AND links
    before sending to the LLM.

    PDF:  pymupdf4llm -> markdown (preserves headings, bold, lists, tables)
          PyMuPDF     -> hyperlink extraction from annotations
    DOCX: mammoth     -> HTML (preserves structure and links)
          BeautifulSoup -> link extraction from HTML
          python-docx -> fallback hyperlink extraction from OOXML rels
    """

    def __init__(self):
        from app.services.ai_service import ai_service
        self.ai_service = ai_service

    async def parse_file(self, file_data: bytes, file_type: str) -> Dict[str, Any]:
        """
        Parse a resume file using hybrid extraction:
        1. Extract structured markdown + raw text
        2. Extract all hyperlinks separately
        3. Categorize links (linkedin, github, etc.)
        4. Send markdown + pre-extracted links to LLM

        Returns:
            Parsed resume data with links preserved in user_details
        """
        if file_type.lower() == "pdf":
            markdown_text, raw_text, links = await self._extract_from_pdf(file_data)
        elif file_type.lower() in ["docx", "doc"]:
            markdown_text, raw_text, links = await self._extract_from_docx(file_data)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Also extract URLs/emails/phones from the raw text as fallback
        text_urls = _extract_urls_from_text(raw_text)
        text_emails = _extract_emails_from_text(raw_text)
        text_phones = _extract_phones_from_text(raw_text)

        # Merge annotation/hyperlink URLs with text-extracted URLs (dedup)
        all_urls = list(set(links + text_urls))
        categorized_links = _categorize_links(all_urls)

        logger.info(
            f"Extracted {len(markdown_text)} chars markdown, "
            f"{len(all_urls)} URLs, {len(text_emails)} emails, "
            f"{len(text_phones)} phones from {file_type} file"
        )
        link_summary = {k: len(v) for k, v in categorized_links.items()}
        logger.info("Categorized links: %s", link_summary)

        # Build pre-extracted contact info to pass alongside markdown
        pre_extracted = {
            "emails": text_emails,
            "phones": text_phones,
            "links": categorized_links,
        }

        # Send structured markdown + pre-extracted info to LLM
        parsed_data = await self.ai_service.extract_resume_data(
            markdown_text, pre_extracted
        )

        # ── Ground-truth overrides ──────────────────────────────────────
        # Our regex extraction is more reliable than LLM for contact info
        # and links, so we overwrite/merge authoritatively.

        # Links
        existing_links = parsed_data.get("extracted_links") or {}
        for cat, urls in categorized_links.items():
            existing = existing_links.get(cat, [])
            existing_links[cat] = list(set(existing + urls))
        parsed_data["extracted_links"] = existing_links

        # Contact fields
        if not parsed_data.get("email") and text_emails:
            parsed_data["email"] = text_emails[0]
        if not parsed_data.get("phone") and text_phones:
            parsed_data["phone"] = text_phones[0]
        if not parsed_data.get("linkedin") and categorized_links.get("linkedin"):
            parsed_data["linkedin"] = categorized_links["linkedin"][0]
        if not parsed_data.get("github") and categorized_links.get("github"):
            parsed_data["github"] = categorized_links["github"][0]
        if not parsed_data.get("portfolio") and categorized_links.get("portfolio"):
            parsed_data["portfolio"] = categorized_links["portfolio"][0]

        parsed_data["raw_text"] = raw_text

        logger.info(
            f"AI extraction completed. Name: {parsed_data.get('name')}, "
            f"Skills: {len(parsed_data.get('skills', []))}, "
            f"Experience: {len(parsed_data.get('experience', []))}, "
            f"Education: {len(parsed_data.get('education', []))}, "
            f"Links preserved: {sum(len(v) for v in categorized_links.values())}"
        )

        return parsed_data

    # ── PDF Extraction ───────────────────────────────────────────────

    async def _extract_from_pdf(self, file_data: bytes) -> Tuple[str, str, List[str]]:
        """
        Extract from PDF using PyMuPDF ecosystem.

        Returns:
            (markdown_text, raw_text, list_of_hyperlink_urls)
        """
        links: List[str] = []
        raw_text_parts: List[str] = []

        try:
            doc = fitz.open(stream=file_data, filetype="pdf")

            # 1) Extract markdown using pymupdf4llm (preserves structure)
            markdown_text = pymupdf4llm.to_markdown(doc)

            # 2) Extract raw text and hyperlinks from annotations
            for page in doc:
                # Raw text for fallback extraction
                page_text = page.get_text("text")
                if page_text:
                    raw_text_parts.append(page_text)

                # Extract hyperlinks from PDF annotations
                for link in page.get_links():
                    uri = link.get("uri", "")
                    if uri and uri.startswith("http"):
                        links.append(uri)

            doc.close()

        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            raise

        raw_text = "\n".join(raw_text_parts)
        return markdown_text, raw_text, links

    # ── DOCX Extraction ──────────────────────────────────────────────

    async def _extract_from_docx(self, file_data: bytes) -> Tuple[str, str, List[str]]:
        """
        Extract from DOCX using mammoth + python-docx.

        Returns:
            (markdown_text, raw_text, list_of_hyperlink_urls)
        """
        links: List[str] = []
        raw_text_parts: List[str] = []

        try:
            file_stream = io.BytesIO(file_data)

            # 1) Convert to HTML using mammoth (preserves structure + links)
            result = mammoth.convert_to_html(file_stream)
            html_content = result.value

            # 2) Extract links from HTML using BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")
            for a_tag in soup.find_all("a", href=True):
                href = a_tag["href"]
                if href and href.startswith("http"):
                    links.append(href)

            # 3) Convert HTML to clean markdown-like text
            markdown_text = self._html_to_markdown(soup)

            # 4) Also extract raw text using python-docx for fallback
            file_stream.seek(0)
            doc = Document(file_stream)

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    raw_text_parts.append(paragraph.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        raw_text_parts.append(" | ".join(row_text))

            # 5) Extract hyperlinks from OOXML relationships (catches links
            #    that mammoth might miss, e.g., field codes)
            file_stream.seek(0)
            docx_links = self._extract_docx_hyperlinks(file_stream)
            links.extend(docx_links)

        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            raise

        raw_text = "\n".join(raw_text_parts)
        links = list(set(links))  # Deduplicate
        return markdown_text, raw_text, links

    def _html_to_markdown(self, soup: BeautifulSoup) -> str:
        """
        Convert BeautifulSoup-parsed HTML to a markdown-like format
        that preserves headings, lists, bold, and links.
        """
        lines: List[str] = []

        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table", "tr"]):
            tag = element.name

            if tag.startswith("h"):
                level = int(tag[1])
                prefix = "#" * level
                text = element.get_text(strip=True)
                if text:
                    lines.append(f"\n{prefix} {text}\n")

            elif tag == "li":
                text = self._element_to_markdown_text(element)
                if text:
                    lines.append(f"- {text}")

            elif tag == "p":
                text = self._element_to_markdown_text(element)
                if text:
                    lines.append(text)

            elif tag == "tr":
                cells = element.find_all(["td", "th"])
                row_text = " | ".join(c.get_text(strip=True) for c in cells)
                if row_text.strip():
                    lines.append(f"| {row_text} |")

        return "\n".join(lines)

    def _element_to_markdown_text(self, element) -> str:
        """Convert an HTML element to markdown text, preserving links and bold."""
        parts: List[str] = []
        for child in element.children:
            if isinstance(child, str):
                parts.append(child)
            elif child.name == "a" and child.get("href"):
                href = child["href"]
                text = child.get_text(strip=True)
                if href.startswith("http"):
                    parts.append(f"[{text}]({href})")
                elif href.startswith("mailto:"):
                    parts.append(href.replace("mailto:", ""))
                else:
                    parts.append(text)
            elif child.name in ("strong", "b"):
                parts.append(f"**{child.get_text(strip=True)}**")
            elif child.name in ("em", "i"):
                parts.append(f"*{child.get_text(strip=True)}*")
            else:
                parts.append(child.get_text(strip=True))
        return " ".join(parts).strip()

    def _extract_docx_hyperlinks(self, file_stream: io.BytesIO) -> List[str]:
        """
        Extract hyperlinks directly from DOCX OOXML relationships.
        This catches links embedded as relationship targets that mammoth may miss.
        """
        links: List[str] = []
        try:
            doc = Document(file_stream)
            # Access the relationships in the main document part
            rels = doc.part.rels
            for rel_id, rel in rels.items():
                if "hyperlink" in str(rel.reltype).lower():
                    target = str(rel._target)
                    if target.startswith("http"):
                        links.append(target)
        except Exception as e:
            logger.warning(f"Could not extract DOCX hyperlinks from rels: {e}")
        return links


# Singleton instance
resume_parser = ResumeParserService()
